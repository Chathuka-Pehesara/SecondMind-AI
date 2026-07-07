import io
import os
import json
import datetime
import uuid
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy import Session
import google.generativeai as genai
import chromadb

from app.database import get_db
from app.models import Document, Conversation
from app.schemas import DocumentResponse
from app.auth import get_current_user

router = APIRouter(prefix="/chat/conversations/{conversation_id}/documents", tags=["rag"])

# Setup chromaDB persistent directory
CHROMA_DB_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "chroma_db")
os.makedirs(CHROMA_DB_DIR, exist_ok=True)

chroma_client = chromadb.PersistentClient(path=CHROMA_DB_DIR)
# Cosine similatiry for scoring
rag_collection = chroma_client.get_or_create_collection(
    name="secondmind_rag",
    metadata={"hnsw:space": "cosine"}
)

# 1. Text Extractor Functions
def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        page_text= page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    text = ""
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\n"
    return text

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")

# 2. Text Chunking
def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    
    return chunks
    

# 3. Gemini Embeddings
def get_embeddings_batch(texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY environment variable is not configured")

    response = genai.embed_content(
        model="models/embedding-004",
        content=texts,
        task_type="retrieval_document"
    )

    return response["embedding"]

def get_query_embedding(query: str) -> List[float]:
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY environment variable is not configured")

    response = genai.embed_content(
        model="models/embedding-004",
        content=query,
        task_type="retrieval_query"
    )

    return response["embedding"]

# 4. Context Retrival Helper
def retrieve_rag_context(conversation_id: str, query: str, top_k: int = 4):
    try:
        query_embed = get_query_embedding(query)
        results = rag_collection.query(
            query_embeddings=[query_embed],
            n_results=top_k,
            where={"conversation_id": conversation_id }
        )

        if not results or not results['documents'] or not results['document'][0]:
            return [],0.0

        chunks = results['documents'][0]
        metadatas = results['metadatas'][0]
        distances = results['distances'][0] if 'distance' in results else []

        similarities = [max(0.0, 1.0 - dist) for dist in distances]
        avg_similarity = sum(similarities) / len(similarities) if similarities else 0.0

        retrived_context = []
        for text, meta in zip(chunks, metadatas):
            retrived_context.append({
                "text": text,
                "filename": meta.get("filename", "Unknown Document")
            }) 
        
        return retrived_context, avg_similarity

    except Exception as e:
        print(f"Error retrieving RAG context: {e}")
        return [], 0.0 
        
# 5. API Endpoints

@router.post("", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    conversation_id: str,
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    conversation = db.query(Conversation).filter(
        Conversation.id == conversation_id,
        Conversation.user_id == current_user.id
    ).first()
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
        
    filename = file.filename
    file_bytes = await file.read()
    file_size = len(file_bytes)
    
    # Extract text content
    try:
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(file_bytes)
        elif filename.endswith(".txt"):
            text = extract_text_from_txt(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF, DOCX, or TXT.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")
        
    if not text.strip():
        raise HTTPException(status_code=400, detail="Document text extraction resulted in an empty string.")
        
    # Chunk text
    chunks = chunk_text(text)
    
    # Store chunks & embeddings in ChromaDB
    try:
        embeddings = get_embeddings_batch(chunks)
        ids = [f"{conversation_id}_{uuid.uuid4()}" for _ in chunks]
        metadatas = [{"conversation_id": conversation_id, "filename": filename, "chunk_index": i} for i in range(len(chunks))]
        
        rag_collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to write to vector DB: {str(e)}")
        
    # Write to local metadata database
    db_doc = Document(
        conversation_id=conversation_id,
        filename=filename,
        file_size=file_size
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    
    return db_doc
@router.get("", response_model=List[DocumentResponse])
def list_documents(
    conversation_id: str,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    conversation = db.query(Conversation).filter(
        Conversation.id == conversation_id,
        Conversation.user_id == current_user.id
    ).first()
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
        
    return db.query(Document).filter(Document.conversation_id == conversation_id).all()
@router.delete("/{document_id}")
def delete_document(
    conversation_id: str,
    document_id: str,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    conversation = db.query(Conversation).filter(
        Conversation.id == conversation_id,
        Conversation.user_id == current_user.id
    ).first()
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
        
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.conversation_id == conversation_id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    # Delete from ChromaDB
    try:
        rag_collection.delete(
            where={"$and": [{"conversation_id": conversation_id}, {"filename": doc.filename}]}
        )
    except Exception as e:
        print(f"ChromaDB cleanup warning: {e}")
        
    db.delete(doc)
    db.commit()
    return {"message": "Document deleted successfully"}


